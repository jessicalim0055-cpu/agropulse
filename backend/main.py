from contextlib import asynccontextmanager
import io
import json
import os
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Header
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from sqlalchemy import func, case
from datetime import datetime, timedelta
from typing import Optional
from dotenv import load_dotenv
import io
import logging
import threading

load_dotenv()

from database import Article, ArticleSentiment, MarketReport, PulseEmailReport, ProcessedEmailId, PriceEntry, LeftfieldReport, create_tables, get_db, SessionLocal
from news_fetcher import fetch_all_feeds
from analyzer import analyze_article, analyze_report, analyze_parity_email, analyze_leftfield_report, COMMODITIES
import vessel_tracker
import outlook_fetcher
import weather as weather_module

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger(__name__)

_refresh_lock = threading.Lock()
_last_refresh: Optional[datetime] = None
_is_refreshing = False


# ── Core refresh logic ────────────────────────────────────────────────────────

def run_refresh(db: Session) -> dict:
    global _last_refresh, _is_refreshing

    if not _refresh_lock.acquire(blocking=False):
        logger.info("Refresh already running — skipping")
        return {"message": "already_running", "new_articles": 0}

    _is_refreshing = True
    try:
        logger.info("Starting news refresh …")
        feeds_data = fetch_all_feeds()
        new_count = 0

        for art in feeds_data:
            if db.query(Article).filter(Article.url == art["url"]).first():
                continue

            article = Article(
                url=art["url"],
                title=art["title"],
                source=art["source"],
                published_at=art["published_at"],
                content=art["content"],
            )
            db.add(article)
            db.flush()  # get article.id before commit

            try:
                result = analyze_article(art["title"], art["content"])
                article.summary = result.get("summary", "")

                for sent in result.get("sentiments", []):
                    db.add(ArticleSentiment(
                        article_id=article.id,
                        commodity=sent["commodity"],
                        sentiment=sent["sentiment"],
                        confidence=sent.get("confidence", 0.7),
                        reasoning=sent.get("reasoning", ""),
                    ))

                article.analyzed = True
            except RuntimeError as e:
                # API key missing — abort analysis but keep article
                logger.error(str(e))
                article.analyzed = False

            new_count += 1

        db.commit()
        _last_refresh = datetime.utcnow()
        logger.info(f"Refresh done — {new_count} new articles")
        return {"message": "ok", "new_articles": new_count}

    except Exception as e:
        db.rollback()
        logger.error(f"Refresh error: {e}")
        raise
    finally:
        _is_refreshing = False
        _refresh_lock.release()


def run_refresh_standalone():
    db = SessionLocal()
    try:
        run_refresh(db)
    except Exception as e:
        logger.error(f"Standalone refresh error: {e}")
    finally:
        db.close()


# ── App lifecycle ─────────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    create_tables()
    # Schedule automatic refresh every 6 hours
    from apscheduler.schedulers.background import BackgroundScheduler
    scheduler = BackgroundScheduler()
    scheduler.add_job(run_refresh_standalone, trigger="interval", hours=6, id="auto_refresh")
    scheduler.add_job(sync_parity_emails_job, trigger="interval", weeks=1, id="parity_email_sync")
    scheduler.add_job(sync_leftfield_job, trigger="interval", weeks=1, id="leftfield_sync")
    scheduler.start()

    # Initial refresh on startup (non-blocking)
    t = threading.Thread(target=run_refresh_standalone, daemon=True)
    t.start()

    # Start live AIS vessel tracking stream
    await vessel_tracker.start()

    yield
    scheduler.shutdown(wait=False)
    vessel_tracker.stop()


app = FastAPI(title="AgroPulse API", version="1.0.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Routes ────────────────────────────────────────────────────────────────────

_news_cache: dict = {}   # {query: (timestamp, articles)}
_NEWS_CACHE_TTL = 3600   # 1 hour

@app.get("/api/conflict-news")
def get_conflict_news(q: str, db: Session = Depends(get_db)):
    import time as _time

    cache_key = q.lower().strip()
    if cache_key in _news_cache:
        ts, cached = _news_cache[cache_key]
        if _time.time() - ts < _NEWS_CACHE_TTL:
            return {"articles": cached}

    articles: list[dict] = []

    # ── 1. Search local DB (title-only for relevance) ─────────────────────────
    terms = [t.strip() for t in q.split() if len(t.strip()) > 2]
    if terms:
        from sqlalchemy import or_
        title_conditions = [Article.title.ilike(f"%{t}%") for t in terms]
        rows = (
            db.query(Article)
            .filter(or_(*title_conditions))
            .order_by(Article.published_at.desc())
            .limit(6)
            .all()
        )
        for a in rows:
            snippet = (a.content or "").strip()[:280].rsplit(" ", 1)[0] + "…" if a.content else ""
            articles.append({
                "title": a.title,
                "description": snippet,
                "url": a.url,
                "image": "",
                "source": a.source or "",
                "published_at": a.published_at.isoformat() if a.published_at else "",
            })

    # ── 2. Top up with NewsAPI if key present ─────────────────────────────────
    api_key = os.getenv("NEWS_API_KEY", "")
    if api_key and not api_key.startswith("your-") and len(articles) < 4:
        try:
            resp = requests.get(
                "https://newsapi.org/v2/everything",
                params={"q": q, "sortBy": "publishedAt", "pageSize": 4, "language": "en", "apiKey": api_key},
                timeout=10,
            )
            if resp.ok:
                existing_urls = {a["url"] for a in articles}
                for a in resp.json().get("articles", []):
                    if a.get("title") and "[Removed]" not in a.get("title", "") and a["url"] not in existing_urls:
                        articles.append({
                            "title": a["title"],
                            "description": a.get("description") or "",
                            "url": a["url"],
                            "image": a.get("urlToImage") or "",
                            "source": a.get("source", {}).get("name", ""),
                            "published_at": a.get("publishedAt", ""),
                        })
        except Exception as e:
            logger.error(f"Conflict news NewsAPI error: {e}")

    articles = articles[:6]
    _news_cache[cache_key] = (_time.time(), articles)
    return {"articles": articles}

@app.get("/api/status")
def get_status(db: Session = Depends(get_db)):
    total = db.query(func.count(Article.id)).scalar() or 0
    cutoff = datetime.utcnow() - timedelta(hours=24)
    recent = db.query(func.count(Article.id)).filter(Article.fetched_at >= cutoff).scalar() or 0
    return {
        "last_refresh": _last_refresh.isoformat() if _last_refresh else None,
        "is_refreshing": _is_refreshing,
        "total_articles": total,
        "articles_last_24h": recent,
    }


@app.post("/api/refresh")
def trigger_refresh():
    if _is_refreshing:
        return {"message": "already_running"}
    t = threading.Thread(target=run_refresh_standalone, daemon=True)
    t.start()
    return {"message": "started"}


@app.get("/api/news")
def get_news(
    limit: int = 30,
    offset: int = 0,
    commodity: Optional[str] = None,
    db: Session = Depends(get_db),
):
    query = db.query(Article).filter(Article.analyzed == True)

    if commodity and commodity in COMMODITIES:
        sub = db.query(ArticleSentiment.article_id).filter(
            ArticleSentiment.commodity == commodity
        ).subquery()
        query = query.filter(Article.id.in_(sub))

    articles = (
        query
        .order_by(Article.published_at.desc())
        .offset(offset)
        .limit(limit)
        .all()
    )

    result = []
    for a in articles:
        sents = db.query(ArticleSentiment).filter(ArticleSentiment.article_id == a.id).all()
        result.append({
            "id": a.id,
            "title": a.title,
            "url": a.url,
            "source": a.source,
            "published_at": a.published_at.isoformat() if a.published_at else None,
            "summary": a.summary,
            "commodities": [
                {
                    "commodity": s.commodity,
                    "name": COMMODITIES.get(s.commodity, s.commodity),
                    "sentiment": s.sentiment,
                    "confidence": s.confidence,
                    "reasoning": s.reasoning,
                }
                for s in sents
            ],
        })
    return result


@app.get("/api/sentiment/current")
def get_current_sentiment(db: Session = Depends(get_db)):
    cutoff = datetime.utcnow() - timedelta(days=7)
    result = {}

    for key, name in COMMODITIES.items():
        rows = (
            db.query(
                func.sum(case((ArticleSentiment.sentiment == "bullish", 1), else_=0)).label("bull"),
                func.sum(case((ArticleSentiment.sentiment == "bearish", 1), else_=0)).label("bear"),
                func.sum(case((ArticleSentiment.sentiment == "neutral", 1), else_=0)).label("neut"),
                func.count().label("total"),
            )
            .join(Article, Article.id == ArticleSentiment.article_id)
            .filter(ArticleSentiment.commodity == key, Article.published_at >= cutoff)
            .one()
        )

        bull, bear, neut, total = (rows.bull or 0), (rows.bear or 0), (rows.neut or 0), (rows.total or 0)
        dominant = max([("bullish", bull), ("bearish", bear), ("neutral", neut)], key=lambda x: x[1])[0] if total else "neutral"

        # Fetch up to 3 recent reasoning snippets for the dominant sentiment
        reasons_rows = (
            db.query(ArticleSentiment.reasoning)
            .join(Article, Article.id == ArticleSentiment.article_id)
            .filter(
                ArticleSentiment.commodity == key,
                ArticleSentiment.sentiment == dominant,
                Article.published_at >= cutoff,
                ArticleSentiment.reasoning != None,
                ArticleSentiment.reasoning != "",
            )
            .order_by(Article.published_at.desc())
            .limit(3)
            .all()
        )
        top_reasons = [r.reasoning[:220] for r in reasons_rows if r.reasoning]

        result[key] = {
            "key": key,
            "name": name,
            "dominant_sentiment": dominant,
            "bullish_count": bull,
            "bearish_count": bear,
            "neutral_count": neut,
            "total_articles": total,
            "bullish_pct": round(bull / total * 100) if total else 0,
            "bearish_pct": round(bear / total * 100) if total else 0,
            "neutral_pct": round(neut / total * 100) if total else 0,
            "net_score": round((bull - bear) / total, 3) if total else 0,
            "top_reasons": top_reasons,
        }

    return result


@app.get("/api/trends")
def get_trends(
    period: str = "week",
    commodity: Optional[str] = None,
    db: Session = Depends(get_db),
):
    delta_map = {"day": 1, "week": 7, "month": 30, "year": 365}
    cutoff = datetime.utcnow() - timedelta(days=delta_map.get(period, 7))

    keys = [commodity] if commodity and commodity in COMMODITIES else list(COMMODITIES.keys())
    data = {}

    for key in keys:
        rows = (
            db.query(
                func.date(Article.published_at).label("date"),
                func.sum(case((ArticleSentiment.sentiment == "bullish", 1), else_=0)).label("bull"),
                func.sum(case((ArticleSentiment.sentiment == "bearish", 1), else_=0)).label("bear"),
                func.sum(case((ArticleSentiment.sentiment == "neutral", 1), else_=0)).label("neut"),
            )
            .join(Article, Article.id == ArticleSentiment.article_id)
            .filter(ArticleSentiment.commodity == key, Article.published_at >= cutoff)
            .group_by(func.date(Article.published_at))
            .order_by(func.date(Article.published_at))
            .all()
        )

        data[key] = {
            "name": COMMODITIES[key],
            "points": [
                {
                    "date": str(r.date),
                    "bullish": r.bull or 0,
                    "bearish": r.bear or 0,
                    "neutral": r.neut or 0,
                    "net_score": round(
                        ((r.bull or 0) - (r.bear or 0)) / max((r.bull or 0) + (r.bear or 0) + (r.neut or 0), 1),
                        3,
                    ),
                }
                for r in rows
            ],
        }

    return data


def _check_admin(password: Optional[str]):
    expected = os.getenv("UPLOAD_PASSWORD", "")
    if expected and password != expected:
        raise HTTPException(status_code=401, detail="Invalid admin password.")


@app.get("/api/reports")
def get_reports(db: Session = Depends(get_db)):
    rows = db.query(MarketReport).order_by(MarketReport.analyzed_at.desc()).all()
    return [{"id": r.id, "analyzed_at": r.analyzed_at.isoformat(), **json.loads(r.data_json)} for r in rows]


@app.post("/api/reports/analyze")
async def analyze_report_endpoint(
    file: UploadFile = File(...),
    x_admin_password: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    _check_admin(x_admin_password)

    filename = file.filename or "report"
    content = await file.read()
    text = ""

    if filename.lower().endswith(".pdf"):
        # Try three engines in order — pymupdf > pdfplumber > pypdf
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
        except Exception:
            pass
        if not text.strip():
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            except Exception:
                pass
        if not text.strip():
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception as e:
                raise HTTPException(status_code=422, detail=f"Could not read PDF: {e}")
    elif filename.lower().endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Could not read DOCX: {e}")
    elif filename.lower().endswith(".txt"):
        text = content.decode("utf-8", errors="replace")
    else:
        raise HTTPException(status_code=422, detail="Unsupported file type. Please upload PDF, DOCX, or TXT.")

    if not text.strip():
        raise HTTPException(status_code=422, detail="No text found in this file. If it is a scanned PDF, try saving it as a Word doc or copying the text into a .txt file first.")

    result = analyze_report(text)
    result["filename"] = filename
    result["analyzed_at"] = datetime.utcnow().isoformat()

    db.add(MarketReport(filename=filename, data_json=json.dumps(result)))
    db.commit()

    return result


@app.delete("/api/reports/{report_id}")
def delete_report(
    report_id: int,
    x_admin_password: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    _check_admin(x_admin_password)
    row = db.query(MarketReport).filter(MarketReport.id == report_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Report not found.")
    db.delete(row)
    db.commit()
    return {"ok": True}


# ── Outlook auto-sync ────────────────────────────────────────────────────────

def sync_parity_emails_job():
    """Called by APScheduler weekly and on-demand. Fetches + analyses new parity emails."""
    db = SessionLocal()
    try:
        seen_ids = {r.message_id for r in db.query(ProcessedEmailId.message_id).all()}
        emails = outlook_fetcher.fetch_parity_emails(already_seen_ids=seen_ids)
        logger.info(f"Outlook sync: {len(emails)} new parity email(s) found")
        for em in emails:
            result = analyze_parity_email(
                em["body"],
                sender=f"{em['from_name']} <{em['from_email']}>",
                email_date=em["received"][:10] if em["received"] else "",
            )
            result.setdefault("sender", f"{em['from_name']} <{em['from_email']}>")
            result.setdefault("email_date", em["received"][:10] if em["received"] else "")
            row = PulseEmailReport(
                filename=em["subject"] or "Outlook email",
                sender=result.get("sender", ""),
                email_date=result.get("email_date", ""),
                data_json=json.dumps({k: v for k, v in result.items() if k not in ("sender", "email_date")}),
            )
            db.add(row)
            db.add(ProcessedEmailId(message_id=em["id"]))
        db.commit()
    except Exception as e:
        db.rollback()
        logger.error(f"Outlook sync error: {e}")
    finally:
        db.close()


@app.get("/api/outlook/status")
def outlook_status():
    status = outlook_fetcher.get_auth_status()
    config = outlook_fetcher.get_config()
    return {**status, **config}


@app.post("/api/outlook/auth/start")
def outlook_auth_start(x_admin_password: Optional[str] = Header(None)):
    _check_admin(x_admin_password)
    if not outlook_fetcher.get_config()["configured"]:
        raise HTTPException(status_code=400, detail="AZURE_CLIENT_ID not set in .env — see setup instructions.")
    try:
        return outlook_fetcher.start_device_flow()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/outlook/auth/status")
def outlook_auth_status():
    return outlook_fetcher.get_auth_status()


@app.post("/api/outlook/disconnect")
def outlook_disconnect(x_admin_password: Optional[str] = Header(None)):
    _check_admin(x_admin_password)
    outlook_fetcher.disconnect()
    return {"ok": True}


@app.post("/api/outlook/sync")
def outlook_sync_now(x_admin_password: Optional[str] = Header(None)):
    _check_admin(x_admin_password)
    if not outlook_fetcher.is_authenticated():
        raise HTTPException(status_code=400, detail="Outlook not authenticated.")
    t = threading.Thread(target=sync_parity_emails_job, daemon=True)
    t.start()
    return {"message": "Sync started"}


# ── Parity email reports ──────────────────────────────────────────────────────

@app.get("/api/parity-emails")
def get_parity_emails(db: Session = Depends(get_db)):
    rows = db.query(PulseEmailReport).order_by(PulseEmailReport.uploaded_at.desc()).all()
    return [
        {"id": r.id, "filename": r.filename, "sender": r.sender,
         "email_date": r.email_date, "uploaded_at": r.uploaded_at.isoformat(),
         **json.loads(r.data_json)}
        for r in rows
    ]


@app.post("/api/parity-emails/upload")
async def upload_parity_email(
    file: UploadFile = File(...),
    x_admin_password: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    _check_admin(x_admin_password)
    filename = file.filename or "email"
    content = await file.read()
    text = ""
    sender = ""
    email_date = ""

    if filename.lower().endswith(".eml"):
        import email as email_lib
        msg = email_lib.message_from_bytes(content)
        sender = msg.get("From", "")
        email_date = msg.get("Date", "")
        # Extract text body
        if msg.is_multipart():
            for part in msg.walk():
                ct = part.get_content_type()
                if ct == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        text += payload.decode(part.get_content_charset() or "utf-8", errors="replace") + "\n"
                elif ct == "text/html" and not text:
                    payload = part.get_payload(decode=True)
                    if payload:
                        import re
                        html = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
                        text += re.sub(r"<[^>]+>", " ", html) + "\n"
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                text = payload.decode(msg.get_content_charset() or "utf-8", errors="replace")
                if msg.get_content_type() == "text/html":
                    import re
                    text = re.sub(r"<[^>]+>", " ", text)
    elif filename.lower().endswith(".txt"):
        text = content.decode("utf-8", errors="replace")
    elif filename.lower().endswith(".pdf"):
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
        except Exception:
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            except Exception as e:
                raise HTTPException(status_code=422, detail=f"Could not read PDF: {e}")
    elif filename.lower().endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Could not read DOCX: {e}")
    else:
        raise HTTPException(status_code=422, detail="Unsupported file type. Please upload .eml, .txt, .pdf, or .docx.")

    if not text.strip():
        raise HTTPException(status_code=422, detail="No text content found in this file.")

    result = analyze_parity_email(text, sender=sender, email_date=email_date)
    result.setdefault("sender", sender)
    result.setdefault("email_date", email_date)

    row = PulseEmailReport(
        filename=filename,
        sender=result.get("sender") or sender,
        email_date=result.get("email_date") or email_date,
        data_json=json.dumps({k: v for k, v in result.items() if k not in ("sender", "email_date")}),
    )
    db.add(row)
    db.commit()
    db.refresh(row)

    return {
        "id": row.id, "filename": row.filename, "sender": row.sender,
        "email_date": row.email_date, "uploaded_at": row.uploaded_at.isoformat(),
        **json.loads(row.data_json),
    }


@app.delete("/api/parity-emails/{report_id}")
def delete_parity_email(
    report_id: int,
    x_admin_password: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    _check_admin(x_admin_password)
    row = db.query(PulseEmailReport).filter(PulseEmailReport.id == report_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Email report not found.")
    db.delete(row)
    db.commit()
    return {"ok": True}


# ── Leftfield Research Reports ───────────────────────────────────────────────

def _extract_pdf_text(pdf_bytes: bytes) -> str:
    """Try multiple PDF engines in order of preference."""
    text = ""
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
    except Exception:
        pass
    if not text.strip():
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception:
            pass
    if not text.strip():
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            pass
    return text


def sync_leftfield_job():
    """Fetch and analyse new research report PDFs from leftfield@leftfieldcr.com."""
    db = SessionLocal()
    try:
        seen_ids = {r.message_id for r in db.query(ProcessedEmailId).filter(
            ProcessedEmailId.message_id.like("leftfield:%")
        ).all()}
        emails = outlook_fetcher.fetch_leftfield_emails(already_seen_ids=seen_ids)
        logger.info(f"Leftfield sync: {len(emails)} new report(s) found")
        for em in emails:
            pdf_text = _extract_pdf_text(em["attachment_bytes"])
            if not pdf_text.strip():
                logger.warning(f"No text extracted from {em['attachment_name']} — skipping")
                continue
            result = analyze_leftfield_report(
                text=pdf_text,
                subject=em.get("subject", ""),
                received_date=em["received"][:10] if em.get("received") else "",
            )
            report_date = result.get("report_date") or (em["received"][:10] if em.get("received") else "")
            row = LeftfieldReport(
                filename=em["attachment_name"],
                email_subject=em.get("subject", ""),
                report_date=report_date,
                data_json=json.dumps({k: v for k, v in result.items() if k != "report_date"}),
            )
            db.add(row)
            db.add(ProcessedEmailId(message_id=em["id"]))
        db.commit()
    except Exception as e:
        db.rollback()
        logger.error(f"Leftfield sync error: {e}")
    finally:
        db.close()


@app.get("/api/leftfield-reports")
def get_leftfield_reports(db: Session = Depends(get_db)):
    rows = db.query(LeftfieldReport).order_by(LeftfieldReport.report_date.desc(), LeftfieldReport.synced_at.desc()).all()
    return [
        {
            "id": r.id,
            "filename": r.filename,
            "email_subject": r.email_subject,
            "report_date": r.report_date,
            "synced_at": r.synced_at.isoformat(),
            **json.loads(r.data_json),
        }
        for r in rows
    ]


@app.post("/api/leftfield/sync")
def leftfield_sync_now(x_admin_password: Optional[str] = Header(None)):
    _check_admin(x_admin_password)
    if not outlook_fetcher.is_authenticated():
        raise HTTPException(status_code=400, detail="Outlook not authenticated.")
    t = threading.Thread(target=sync_leftfield_job, daemon=True)
    t.start()
    return {"message": "Sync started"}


@app.delete("/api/leftfield-reports/{report_id}")
def delete_leftfield_report(
    report_id: int,
    x_admin_password: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    _check_admin(x_admin_password)
    row = db.query(LeftfieldReport).filter(LeftfieldReport.id == report_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Report not found.")
    db.delete(row)
    db.commit()
    return {"ok": True}


# ── Weather ───────────────────────────────────────────────────────────────────

@app.get("/api/weather")
def get_weather():
    """Live agricultural weather for all tracked regions (cached 3h)."""
    try:
        return weather_module.fetch_all_weather()
    except Exception as e:
        logger.error(f"Weather endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/weather/refresh")
def refresh_weather():
    """Force-clear the weather cache so the next GET re-fetches."""
    weather_module._cache.clear()
    return {"message": "Cache cleared"}


# ── Price tracking ────────────────────────────────────────────────────────────

from pydantic import BaseModel


class PriceEntryCreate(BaseModel):
    commodity: str
    origin: str
    destination: str
    price: float
    trade_type: str   # buy | sell | indicative
    cargo_type: str   # bulk | container
    date: str         # YYYY-MM-DD
    notes: Optional[str] = None


def _price_row(r: PriceEntry) -> dict:
    return {
        "id": r.id,
        "commodity": r.commodity,
        "origin": r.origin,
        "destination": r.destination,
        "price": r.price,
        "trade_type": r.trade_type,
        "cargo_type": r.cargo_type,
        "date": r.date,
        "notes": r.notes,
    }


@app.get("/api/prices")
def get_prices(
    commodity: Optional[str] = None,
    trade_type: Optional[str] = None,
    cargo_type: Optional[str] = None,
    db: Session = Depends(get_db),
):
    q = db.query(PriceEntry).order_by(PriceEntry.date.desc(), PriceEntry.created_at.desc())
    if commodity:
        q = q.filter(PriceEntry.commodity == commodity)
    if trade_type:
        q = q.filter(PriceEntry.trade_type == trade_type)
    if cargo_type:
        q = q.filter(PriceEntry.cargo_type == cargo_type)
    return [_price_row(r) for r in q.all()]


@app.post("/api/prices")
def create_price(
    body: PriceEntryCreate,
    x_admin_password: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    _check_admin(x_admin_password)
    if body.commodity not in COMMODITIES:
        raise HTTPException(status_code=422, detail="Unknown commodity key.")
    if body.trade_type not in ("buy", "sell", "indicative"):
        raise HTTPException(status_code=422, detail="trade_type must be buy, sell, or indicative.")
    if body.cargo_type not in ("bulk", "container"):
        raise HTTPException(status_code=422, detail="cargo_type must be bulk or container.")
    entry = PriceEntry(
        commodity=body.commodity,
        origin=body.origin.strip(),
        destination=body.destination.strip(),
        price=body.price,
        trade_type=body.trade_type,
        cargo_type=body.cargo_type,
        date=body.date,
        notes=body.notes,
    )
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return _price_row(entry)


@app.delete("/api/prices/{entry_id}")
def delete_price(
    entry_id: int,
    x_admin_password: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    _check_admin(x_admin_password)
    row = db.query(PriceEntry).filter(PriceEntry.id == entry_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Price entry not found.")
    db.delete(row)
    db.commit()
    return {"ok": True}


# ── Vessel tracking ───────────────────────────────────────────────────────────



class VesselTrackRequest(BaseModel):
    imo: list[str]
    mmsi_map: dict[str, str] = {}   # { imo -> mmsi }

@app.post("/api/vessel-positions")
def post_vessel_positions(body: VesselTrackRequest):
    """Register vessels to track (by IMO + MMSI) and return current cached positions."""
    vessel_tracker.update_tracked(body.imo, body.mmsi_map)
    return vessel_tracker.get_positions(body.imo)


@app.get("/api/vessel-positions")
def get_vessel_positions(imo: str = ""):
    """Return cached live positions for comma-separated IMO numbers."""
    imos = [i.strip() for i in imo.split(",") if i.strip()]
    return vessel_tracker.get_positions(imos)


@app.get("/api/vessels/lookup")
def lookup_vessel(imo: str):
    """Look up static vessel info (name, type, flag) by IMO — called when admin adds a vessel."""
    if not imo or not imo.strip().isdigit():
        raise HTTPException(status_code=400, detail="IMO must be a numeric string.")
    result = vessel_tracker.lookup_vessel(imo.strip())
    return result
